import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from schemas import PDD, ExceptionItem
from flowchart import build_flowchart


def set_cell_shading(cell, color_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def remove_heading_border(paragraph):
    """Remove the bottom border that comes by default with Title/Heading styles."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "none")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_bookmark(paragraph, bookmark_name: str, bookmark_id: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    paragraph._p.insert(0, start)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, bookmark_name: str, text: str):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_manual_toc(doc, entries: list[tuple[str, str, int]]):
    """entries: list of (bookmark_name, display_text, indent_level)"""
    for bookmark_name, text, indent in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25 * indent)
        add_internal_hyperlink(p, bookmark_name, text)


def add_exception_table(doc, title: str, items: list[ExceptionItem]):
    doc.add_heading(title, level=2)

    if not items:
        doc.add_paragraph(f"No {title.lower()} identified in this process.")
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Exception Name", "Action", "Parameters", "Action to be Taken"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "2F5496")

    for item in items:
        row = table.add_row()
        row.cells[0].text = item.name
        row.cells[1].text = item.action
        row.cells[2].text = item.parameters
        row.cells[3].text = item.action_to_be_taken


def render_docx(pdd: PDD, output_path: str = "output/PDD_final.docx") -> str:
    doc = Document()

    # --- Page 1: Cover ---
    title = doc.add_heading("Process Design Document (PDD)", level=0)
    remove_heading_border(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading(pdd.process_name, level=1)
    remove_heading_border(subtitle)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- Page 2: Table of Contents (manual, hyperlinked) ---
    toc_heading = doc.add_heading("Table of Contents", level=1)
    remove_heading_border(toc_heading)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_entries = [
        ("intro", "1. Introduction", 0),
        ("proposal", "1.1 Project Proposal", 1),
        ("asis", "2. As Is", 0),
        ("flowchart_sec", "2.1 Flowchart", 1),
        ("steps_sec", "2.2 Process Steps", 1),
        ("exceptions", "3. Exceptions", 0),
    ]
    add_manual_toc(doc, toc_entries)

    doc.add_page_break()

    # --- Page 3+: Content ---

    # 1. Introduction
    h = doc.add_heading("1. Introduction", level=1)
    add_bookmark(h, "intro", 1)

    h = doc.add_heading("1.1 Project Proposal", level=2)
    add_bookmark(h, "proposal", 2)
    doc.add_paragraph(pdd.project_proposal)

    # 2. As Is
    h = doc.add_heading("2. As Is", level=1)
    add_bookmark(h, "asis", 3)

    h = doc.add_heading("2.1 Flowchart", level=2)
    add_bookmark(h, "flowchart_sec", 4)
    try:
        flowchart_path = build_flowchart(pdd)
        doc.add_picture(flowchart_path, width=Inches(3.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"Failed to generate/insert flowchart: {e}")
        doc.add_paragraph("[Flowchart could not be generated]")

    h = doc.add_heading("2.2 Process Steps", level=2)
    add_bookmark(h, "steps_sec", 5)
    for step in pdd.as_is:
        doc.add_heading(f"Step {step.number}", level=3)

        doc.add_paragraph(step.instruction)

        try:
            doc.add_picture(step.frame_ref, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption = doc.add_paragraph(f"Screenshot — Step {step.number}")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].italic = True
            caption.runs[0].font.size = Pt(9)
        except FileNotFoundError:
            doc.add_paragraph(f"[Image not found: {step.frame_ref}]")

        doc.add_paragraph()

    # 3. Exceptions
    h = doc.add_heading("3. Exceptions", level=1)
    add_bookmark(h, "exceptions", 6)
    add_exception_table(doc, "Business Exceptions", pdd.business_exceptions)
    add_exception_table(doc, "System Exceptions", pdd.system_exceptions)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    return output_path